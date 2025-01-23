import streamlit as st
from openai import OpenAI
import json
import random
import PyPDF2
import docx
import re
import base64
from pdf2image import convert_from_bytes
import io
from PIL import Image
import logging
import streamlit.components.v1 as components
import httpx
import os

# Logging für bessere Fehlerverfolgung einrichten
logging.basicConfig(level=logging.INFO)

# Seitenkonfiguration festlegen und Light Mode erzwingen
st.set_page_config(
    page_title="import streamlit as st
import streamlit.components.v1 as components
from openai import OpenAI
import json
import random
import PyPDF2
import docx
import re
import base64
from pdf2image import convert_from_bytes
import io
from PIL import Image
import logging
import httpx
import os

# Set page title and icon
st.set_page_config(page_title="OLAT Fragen Generator BMS", page_icon="📝", layout="wide", initial_sidebar_state="expanded")

# Set up logging for better error tracking
logging.basicConfig(level=logging.INFO)

# Enforce Light Mode using CSS
st.markdown(
    """
    <style>
    /* Force light mode */
    body, .css-18e3th9, .css-1d391kg {
        background-color: white;
        color: black;
    }
    /* Override Streamlit's default dark mode elements */
    .css-1aumxhk, .css-1v3fvcr {
        background-color: white;
    }
    /* Ensure all text is dark */
    .css-1v0mbdj, .css-1xarl3l {
        color: black;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# Clear any existing proxy environment variables to prevent OpenAI SDK from using them
os.environ.pop('HTTP_PROXY', None)
os.environ.pop('HTTPS_PROXY', None)
os.environ.pop('http_proxy', None)
os.environ.pop('https_proxy', None)

# Initialize a custom httpx client without proxies
http_client = httpx.Client()

# Initialize OpenAI client with Streamlit Secrets
try:
    client = OpenAI(
        api_key=st.secrets["openai"]["api_key"],  # API key from Streamlit Secrets
        http_client=http_client
    )
    st.success("OpenAI client initialized successfully.")
except Exception as e:
    st.error(f"Error initializing OpenAI client: {e}")
    st.stop()

# List of available message types
MESSAGE_TYPES = [
    "single_choice",
    "multiple_choice1",
    "multiple_choice2",
    "multiple_choice3",
    "kprim",
    "truefalse",
    "draganddrop",
    "inline_fib"
]

@st.cache_data
def read_prompt_from_md(filename):
    """Read the prompt from a markdown file and cache the result."""
    with open(f"{filename}.md", "r", encoding="utf-8") as file:
        return file.read()

def process_image(_image):
    """Process and resize an image to reduce memory footprint."""
    if isinstance(_image, (str, bytes)):
        img = Image.open(io.BytesIO(base64.b64decode(_image) if isinstance(_image, str) else _image))
    elif isinstance(_image, Image.Image):
        img = _image
    else:
        img = Image.open(_image)

    # Convert to RGB mode if it's not
    if img.mode != 'RGB':
        img = img.convert('RGB')

    # Resize if the image is too large
    max_size = 1000  # Reduced max size to reduce memory consumption
    if max(img.size) > max_size:
        img.thumbnail((max_size, max_size))

    # Save to bytes
    img_byte_arr = io.BytesIO()
    img.save(img_byte_arr, format='JPEG')
    img_byte_arr = img_byte_arr.getvalue()

    return base64.b64encode(img_byte_arr).decode('utf-8')

def replace_german_sharp_s(text):
    """Replace all occurrences of 'ß' with 'ss'."""
    return text.replace('ß', 'ss')

def clean_json_string(s):
    # Remove all markdown code blocks and surrounding whitespace
    s = re.sub(r'^\s*```(json)?\s*', '', s, flags=re.IGNORECASE | re.MULTILINE)
    s = re.sub(r'\s*```\s*$', '', s, flags=re.IGNORECASE | re.MULTILINE)
    
    # Remove any remaining triple backticks in the content
    s = re.sub(r'```', '', s)
    
    # Additional cleaning
    s = s.strip()
    s = re.sub(r'\s+', ' ', s)
    s = re.sub(r'(?<=text": ")(.+?)(?=")', lambda m: m.group(1).replace('\n', '\\n'), s)
    s = ''.join(char for char in s if ord(char) >= 32 or char == '\n')
    
    # Handle potential incomplete JSON
    if not s.startswith('['):
        s = '[' + s
    if not s.endswith(']'):
        s += ']'
    
    return s

def convert_json_to_text_format(json_input):
    if isinstance(json_input, str):
        data = json.loads(json_input)
    else:
        data = json_input

    fib_output = []
    ic_output = []

    for item in data:
        text = item.get('text', '')
        blanks = item.get('blanks', [])
        wrong_substitutes = item.get('wrong_substitutes', [])

        num_blanks = len(blanks)

        fib_lines = [
            "Type\tFIB",
            "Title\t✏✏Vervollständigen Sie die Lücken mit dem korrekten Begriff.✏✏",
            f"Points\t{num_blanks}"
        ]

        for blank in blanks:
            text = text.replace(blank, "{blank}", 1)

        parts = text.split("{blank}")
        for index, part in enumerate(parts):
            fib_lines.append(f"Text\t{part.strip()}")
            if index < len(blanks):
                fib_lines.append(f"1\t{blanks[index]}\t20")

        fib_output.append('\n'.join(fib_lines))

        ic_lines = [
            "Type\tInlinechoice",
            "Title\tWörter einordnen",
            "Question\t✏✏Wählen Sie die richtigen Wörter.✏✏",
            f"Points\t{num_blanks}"
        ]

        all_options = blanks + wrong_substitutes
        random.shuffle(all_options)

        for index, part in enumerate(parts):
            ic_lines.append(f"Text\t{part.strip()}")
            if index < len(blanks):
                options_str = '|'.join(all_options)
                ic_lines.append(f"1\t{options_str}\t{blanks[index]}\t|")

        ic_output.append('\n'.join(ic_lines))

    return '\n\n'.join(fib_output), '\n\n'.join(ic_output)

def transform_output(json_string):
    try:
        cleaned_json_string = clean_json_string(json_string)
        json_data = json.loads(cleaned_json_string)
        fib_output, ic_output = convert_json_to_text_format(json_data)
        
        # Apply the cleaning function here
        fib_output = replace_german_sharp_s(fib_output)
        ic_output = replace_german_sharp_s(ic_output)

        return f"{ic_output}\n---\n{fib_output}"
    except json.JSONDecodeError as e:
        st.error(f"Error parsing JSON: {e}")
        st.text("Cleaned input:")
        st.code(cleaned_json_string, language='json')
        st.text("Original input:")
        st.code(json_string)
        
        try:
            if not cleaned_json_string.strip().endswith(']'):
                cleaned_json_string += ']'
            partial_json = json.loads(cleaned_json_string)
            st.warning("Attempted to salvage partial JSON. Results may be incomplete.")
            fib_output, ic_output = convert_json_to_text_format(partial_json)
            return f"{ic_output}\n---\n{fib_output}"
        except Exception as e_partial:
            st.error(f"Unable to salvage partial JSON: {e_partial}")
            return "Error: Invalid JSON format"
    except Exception as e:
        st.error(f"Error processing input: {str(e)}")
        st.text("Original input:")
        st.code(json_string)
        return "Error: Unable to process input"

def get_chatgpt_response(prompt, image=None, selected_language="English"):
    """Fetch response from OpenAI GPT with error handling."""
    try:
        # Auto-select model based on input type
        model = "gpt-4o-mini" if image else "gpt-4o"
        
        # Create a system prompt that includes language instruction
        system_prompt = (
            """
            You are an expert educator specializing in generating test questions and answers across all topics, following Bloom’s Taxonomy. Your role is to create high-quality Q&A sets based on the material provided by the user, ensuring each question aligns with a specific level of Bloom’s Taxonomy: Remember, Understand, Apply, Analyze, Evaluate, and Create.

            The user will provide input by either uploading a text or an image. Your tasks are as follows:

            Input Analysis:
            - Carefully analyze the content to understand the key concepts and important information.
            - For Images: Analyze diagrams, charts, or infographics to derive educational content.

            Question Generation by Bloom Level:
            Based on the analyzed material (from text or image), generate questions across all six levels of Bloom’s Taxonomy:
            - Remember: Simple recall-based questions.
            - Understand: Questions that assess comprehension of the material.
            - Apply: Questions requiring the use of knowledge in practical situations.
            """
        )
        
        if image:
            base64_image = process_image(image)
            messages = [
                {"role": "system", "content": system_prompt},
                {
                    "role": "user", 
                    "content": [
                        {"type": "text", "text": prompt},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{base64_image}",
                                "detail": "low"
                            }
                        }
                    ]
                }
            ]
        else:
            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": prompt}
            ]

        response = client.chat.completions.create(
            model=model,
            messages=messages,
            max_tokens=16000,
            temperature=0.6
        )
        
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"Error communicating with OpenAI API: {e}")
        logging.error(f"Error communicating with OpenAI API: {e}")
        return None

def process_images(images, selected_language):
    """Process uploaded images and generate questions."""
    for idx, image in enumerate(images):
        st.image(image, caption=f'Page {idx+1}', use_column_width=True)

        # Text area for user input and learning goals
        user_input = st.text_area(f"Enter your question or instructions for Page {idx+1}:", key=f"text_area_{idx}")
        learning_goals = st.text_area(f"Learning Goals for Page {idx+1} (Optional):", key=f"learning_goals_{idx}")
        selected_types = st.multiselect(f"Select question types for Page {idx+1}:", MESSAGE_TYPES, key=f"selected_types_{idx}")

        # Button to generate questions for the page
        if st.button(f"Generate Questions for Page {idx+1}", key=f"generate_button_{idx}"):
            if user_input and selected_types:
                generate_questions_with_image(user_input, learning_goals, selected_types, image, selected_language)
            else:
                st.warning(f"Please enter text and select question types for Page {idx+1}.")

def generate_questions_with_image(user_input, learning_goals, selected_types, image, selected_language):
    """Generate questions for the image and handle errors."""
    all_responses = ""
    generated_content = {}
    for msg_type in selected_types:
        prompt_template = read_prompt_from_md(msg_type)
        full_prompt = f"{prompt_template}\n\nUser Input: {user_input}\n\nLearning Goals: {learning_goals}"
        try:
            response = get_chatgpt_response(full_prompt, image=image, selected_language=selected_language)
            if response:
                if msg_type == "inline_fib":
                    processed_response = transform_output(response)
                    generated_content[f"{msg_type.replace('_', ' ').title()} (Processed)"] = processed_response
                    all_responses += f"{processed_response}\n\n"
                else:
                    generated_content[msg_type.replace('_', ' ').title()] = response
                    all_responses += f"{response}\n\n"
            else:
                st.error(f"Failed to generate a response for {msg_type}.")
        except Exception as e:
            st.error(f"An error occurred for {msg_type}: {str(e)}")
    
    # Apply cleaning function to all responses
    all_responses = replace_german_sharp_s(all_responses)

    # Display generated content with checkmarks
    st.subheader("Generated Content:")
    for title in generated_content.keys():
        st.write(f"✔ {title}")

    # Download button for all responses
    if all_responses:
        st.download_button(
            label="Download All Responses",
            data=all_responses,
            file_name="all_responses.txt",
            mime="text/plain"
        )

@st.cache_data
def convert_pdf_to_images(file):
    """Convert PDF pages to images."""
    images = convert_from_bytes(file.read())
    return images

@st.cache_data
def extract_text_from_pdf(file):
    """Extract text from PDF using PyPDF2."""
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text
    return text.strip()

@st.cache_data
def extract_text_from_docx(file):
    """Extract text from DOCX file."""
    doc = docx.Document(file)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return text.strip()

def is_pdf_ocr(text):
    """Check if PDF contains OCR text."""
    return bool(text)

def process_pdf(file):
    text_content = extract_text_from_pdf(file)
    
    if not text_content or not is_pdf_ocr(text_content):
        st.warning("This PDF is not OCRed. Text extraction failed. Please upload an OCRed PDF.")
        return None, convert_pdf_to_images(file)
    else:
        return text_content, None

def main():
    """Main function for the Streamlit app."""
    st.title("OLAT Fragen Generator BMS")

    # Two-column layout
    col1, col2 = st.columns([1, 2])

    with col1:
        # Expanders above language selection
        with st.expander("ℹ️ Kosteninformationen"):
            st.markdown('''
            <div class="custom-info">
                <ul>
                    <li>Die Nutzungskosten hängen von der <strong>Länge der Eingabe</strong> ab (zwischen $0,01 und $0,1).</li>
                    <li>Jeder ausgewählte Fragetyp kostet ungefähr $0,01.</li>
                </ul>
            </div>
            ''', unsafe_allow_html=True)
        
        with st.expander("✅ Fragetypen"):
            st.markdown('''
            <div class="custom-success">
                <strong>Multiple-Choice-Fragen:</strong>
                <ul>
                    <li>Alle Multiple-Choice-Fragen haben maximal <strong>3 Punkte</strong>.</li>
                    <li><strong>multiple_choice1</strong>: 1 von 4 richtigen Antworten = 3 Punkte</li>
                    <li><strong>multiple_choice2</strong>: 2 von 4 richtigen Antworten = 3 Punkte</li>
                    <li><strong>multiple_choice3</strong>: 3 von 4 richtigen Antworten = 3 Punkte</li>
                </ul>
                <p>Man kann die Punktzahl der Fragen im Editor später mit Ctrl+H suchen und ersetzen. Achtung: Punktzahl für korrekte Antworten UND maximale Punktzahl anpassen!</p>
            </div>
            ''', unsafe_allow_html=True)
            st.markdown('''
            <div class="custom-success">
                <strong>Inline/FIB-Fragen:</strong>
                <ul>
                    <li>Die <strong>Inline</strong>- und <strong>FiB</strong>-Fragen sind inhaltlich identisch.</li>
                    <li>FiB = Das fehlende Wort eingeben.</li>
                    <li>Inline = Das fehlende Wort auswählen.</li>
                </ul>
            </div>
            ''', unsafe_allow_html=True)
            st.markdown('''
            <div class="custom-success">
                <strong>Andere Fragetypen:</strong>
                <ul>
                    <li><strong>Einzelauswahl</strong>: 4 Antworten, 1 Punkt pro Frage.</li>
                    <li><strong>KPRIM</strong>: 4 Antworten, 5 Punkte (4/4 korrekt), 2,5 Punkte (3/4 korrekt), 0 Punkte (50 % oder weniger korrekt).</li>
                    <li><strong>Wahr/Falsch</strong>: 3 Antworten, 3 Punkte pro Frage.</li>
                    <li><strong>Drag & Drop</strong>: Variable Punkte.</li>
                </ul>
            </div>
            ''', unsafe_allow_html=True)
        
        with st.expander("⚠️ Warnungen"):
            st.markdown('''
            <div class="custom-warning">
                <ul>
                    <li><strong>Überprüfen Sie immer, ob die Gesamtpunktzahl = Summe der Punkte der richtigen Antworten ist.</strong></li>
                    <li><strong>Überprüfen Sie immer den Inhalt der Antworten.</strong></li>
                </ul>
            </div>
            ''', unsafe_allow_html=True)

        with st.expander("📧 Kontaktinformationen"):
            st.markdown('''
            <div class="custom-info">
                <p>Wenn du Fragen oder Verbesserungsideen hast, kannst du mich gerne kontaktieren:</p>
                <ul>
                    <li><strong>Pietro Rossi</strong></li>
                    <li><strong>E-Mail:</strong> pietro.rossi[at]bbw.ch</li>
                </ul>
                <p>Ich freue mich über dein Feedback!</p>
            </div>
            ''', unsafe_allow_html=True)

        # Language selection below expanders
        st.markdown("### Sprache auswählen:")
        languages = {
            "German": "German",
            "English": "English",
            "French": "French",
            "Italian": "Italian",
            "Spanish": "Spanish"
        }
        selected_language = st.radio(
            "Wählen Sie die Sprache für den Output:",
            list(languages.values()),
            index=0
        )

    with col2:
        # Video iframe filling the entire right column
        st.markdown("### Videoanleitung")
        components.html(
            """
            <iframe src="https://bbwch-my.sharepoint.com/personal/pietro_rossi_bbw_ch/_layouts/15/embed.aspx?UniqueId=2536d633-4608-4236-a19f-70595426359f&embed=%7B%22ust%22%3Atrue%2C%22hv%22%3A%22CopyEmbedCode%22%7D&referrer=StreamWebApp&referrerScenario=EmbedDialog.Create" width="640" height="360" frameborder="0" scrolling="no" allowfullscreen title="OLAT Fragengenerator.mp4"></iframe>
            """,
            height=370
        )

    # File uploader section
    uploaded_file = st.file_uploader("Upload a PDF, DOCX, or image file", type=["pdf", "docx", "jpg", "jpeg", "png"])

    text_content = ""
    image_content = None
    images = []

    if uploaded_file:
        st.cache_data.clear()

    if uploaded_file is not None:
        if uploaded_file.type == "application/pdf":
            text_content, images = process_pdf(uploaded_file)
            if text_content:
                st.success("Text extracted from PDF. You can now edit it below.")
            elif images:
                st.success("PDF converted to images. You can now ask questions about each page.")
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text_content = extract_text_from_docx(uploaded_file)
            st.success("Text extracted successfully. You can now edit it below.")
        elif uploaded_file.type.startswith('image/'):
            image_content = Image.open(uploaded_file)
            st.image(image_content, caption='Uploaded Image', use_column_width=True)
            st.success("Image uploaded successfully. You can now ask questions about the image.")
        else:
            st.error("Unsupported file type. Please upload a PDF, DOCX, or image file.")

    if images:
        process_images(images, selected_language)
    else:
        user_input = st.text_area("Enter your text or question about the image:", value=text_content)
        learning_goals = st.text_area("Learning Goals (Optional):")
        selected_types = st.multiselect("Select question types to generate:", MESSAGE_TYPES)

        # Custom CSS for styling
        st.markdown(
            """
            <style>
            /* Adaptive CSS for light/dark modes */
            .custom-info {
                background-color: rgba(33, 150, 243, 0.1);
                padding: 10px;
                border-radius: 5px;
                border-left: 6px solid #2196F3;
                color: inherit;
            }
            
            .custom-success {
                background-color: rgba(40, 167, 69, 0.1);
                padding: 10px;
                border-radius: 5px;
                border-left: 6px solid #28a745;
                color: inherit;
            }
            
            .custom-warning {
                background-color: rgba(255, 193, 7, 0.1);
                padding: 10px;
                border-radius: 5px;
                border-left: 6px solid #ffc107;
                color: inherit;
            }
        
            /* Force text color inheritance */
            .custom-info, .custom-success, .custom-warning,
            .custom-info p, .custom-success p, .custom-warning p,
            .custom-info li, .custom-success li, .custom-warning li {
                color: inherit !important;
            }
        
            /* Better contrast for dark mode */
            @media (prefers-color-scheme: dark) {
                .custom-info {
                    background-color: rgba(33, 150, 243, 0.2);
                }
                .custom-success {
                    background-color: rgba(40, 167, 69, 0.2);
                }
                .custom-warning {
                    background-color: rgba(255, 193, 7, 0.2);
                }
            }
            </style>
            """,
            unsafe_allow_html=True
        )

        if st.button("Generate Questions"):
            if (user_input or image_content) and selected_types:
                generate_questions_with_image(user_input, learning_goals, selected_types, image_content, selected_language)              
            elif not user_input and not image_content:
                st.warning("Please enter some text, upload a file, or upload an image.")
            elif not selected_types:
                st.warning("Please select at least one question type.")

if __name__ == "__main__":
    main()",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Enforce Light Mode using CSS
st.markdown(
    """
    <style>
    /* Force light mode */
    body, .css-18e3th9, .css-1d391kg {
        background-color: white;
        color: black;
    }
    /* Override Streamlit's default dark mode elements */
    .css-1aumxhk, .css-1v3fvcr {
        background-color: white;
    }
    /* Ensure all text is dark */
    .css-1v0mbdj, .css-1xarl3l {
        color: black;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# Titel der App
st.title("📝 OLAT Fragen Generator BMS")

# Seitenleiste für Anweisungen und Zusatzinformationen
with st.sidebar:
    st.header("❗ **So verwenden Sie diese App**")
    
    st.markdown("""
    1. **Laden Sie eine PDF, DOCX oder Bilddatei hoch**: Wählen Sie eine Datei von Ihrem Computer aus.
    """)
    
    # YouTube-Video in die Seitenleiste einbetten
    components.html("""
        <iframe width="100%" height="180" src="https://www.youtube.com/embed/NsTAjBdHb1k" 
        title="Demo-Video auf Deutsch" frameborder="0" allow="accelerometer; autoplay; 
        clipboard-write; encrypted-media; gyroscope; picture-in-picture" allowfullscreen>
        </iframe>
    """, height=180)
    
    # Weitere Anweisungen
    st.markdown("""
    2. **Sprache auswählen**: Wählen Sie die gewünschte Sprache für die generierten Fragen.
    3. **Fragetypen auswählen**: Wählen Sie die Typen der Fragen, die Sie generieren möchten.
    4. **Fragen generieren**: Klicken Sie auf die Schaltfläche "Fragen generieren", um den Prozess zu starten.
    5. **Generierte Inhalte herunterladen**: Nach der Generierung können Sie die Antworten herunterladen.
    """)
    
    # Kosteninformationen und Frage-Erklärungen
    st.markdown('''
    <div class="custom-info">
        <strong>ℹ️ Kosteninformationen:</strong>
        <ul>
            <li>Die Nutzungskosten hängen von der <strong>Länge der Eingabe</strong> ab (zwischen 0,01 $ und 0,1 $).</li>
            <li>Jeder ausgewählte Fragetyp kostet ungefähr 0,01 $.</li>
        </ul>
    </div>
    ''', unsafe_allow_html=True)

    st.markdown('''
    <div class="custom-success">
        <strong>✅ Multiple-Choice-Fragen:</strong>
        <ul>
            <li>Alle Multiple-Choice-Fragen haben maximal 3 Punkte.</li>
            <li><strong>multiple_choice1</strong>: 1 von 4 richtigen Antworten.</li>
            <li><strong>multiple_choice2</strong>: 2 von 4 richtigen Antworten.</li>
            <li><strong>multiple_choice3</strong>: 3 von 4 richtigen Antworten.</li>
        </ul>
    </div>
    ''', unsafe_allow_html=True)

    st.markdown('''
    <div class="custom-success">
        <strong>✅ Inline/FIB-Fragen:</strong>
        <ul>
            <li>Die <strong>Inline</strong> und <strong>FIB</strong> Fragen sind inhaltlich identisch.</li>
            <li>FIB = fehlendes Wort eingeben.</li>
            <li>Inline = fehlendes Wort auswählen.</li>
        </ul>
    </div>
    ''', unsafe_allow_html=True)

    st.markdown('''
    <div class="custom-success">
        <strong>✅ Andere Fragetypen:</strong>
        <ul>
            <li><strong>Single Choice</strong>: 4 Antworten, 1 Punkt pro Frage.</li>
            <li><strong>KPRIM</strong>: 4 Antworten, 5 Punkte (4/4 korrekt), 2,5 Punkte (3/4 korrekt), 0 Punkte (50% oder weniger korrekt).</li>
            <li><strong>True/False</strong>: 3 Antworten, 3 Punkte pro Frage.</li>
            <li><strong>Drag & Drop</strong>: Variable Punkte.</li>
        </ul>
    </div>
    ''', unsafe_allow_html=True)

    st.markdown('''
    <div class="custom-warning">
        <strong>⚠️ Warnungen:</strong>
        <ul>
            <li><strong>Überprüfen Sie immer, dass die Gesamtpunkte = Summe der Punkte der korrekten Antworten sind.</strong></li>
            <li><strong>Überprüfen Sie immer den Inhalt der Antworten.</strong></li>
        </ul>
    </div>
    ''', unsafe_allow_html=True)

    # Trennlinie und Lizenzinformationen
    st.markdown("---")
    st.header("📜 Lizenz")
    st.markdown("""
    Diese Anwendung steht unter der [MIT-Lizenz](https://opensource.org/licenses/MIT). 
    Sie dürfen diese Software verwenden, ändern und weitergeben, solange die ursprüngliche Lizenz beibehalten wird.
    """)

    # Kontaktinformationen
    st.header("💬 Kontakt")
    st.markdown("""
    Für Unterstützung, Fragen oder um mehr über die Nutzung dieser App zu erfahren, kannst du gerne auf mich zukommen.
    **Kontakt**: [Pietro](mailto:pietro.rossi@bbw.ch)
    """)

# Clear any existing proxy environment variables to prevent OpenAI SDK from using them
os.environ.pop('HTTP_PROXY', None)
os.environ.pop('HTTPS_PROXY', None)
os.environ.pop('http_proxy', None)
os.environ.pop('https_proxy', None)

# Initialize a custom httpx client without proxies
http_client = httpx.Client()

# Initialize OpenAI client with Streamlit Secrets
try:
    client = OpenAI(
        api_key=st.secrets["openai"]["api_key"],
        http_client=http_client
    )
    st.success("OpenAI-Client erfolgreich initialisiert.")
except Exception as e:
    st.error(f"Fehler bei der Initialisierung des OpenAI-Clients: {e}")
    st.stop()

# Liste der verfügbaren Fragetypen
MESSAGE_TYPES = [
    "single_choice",
    "multiple_choice1",
    "multiple_choice2",
    "multiple_choice3",
    "kprim",
    "truefalse",
    "draganddrop",
    "inline_fib"
]

@st.cache_data
def read_prompt_from_md(filename):
    """Liest den Prompt aus einer Markdown-Datei und speichert das Ergebnis zwischen."""
    with open(f"{filename}.md", "r", encoding="utf-8") as file:
        return file.read()

def process_image(_image):
    """Verarbeitet und verkleinert ein Bild, um den Speicherverbrauch zu reduzieren."""
    if isinstance(_image, (str, bytes)):
        img = Image.open(io.BytesIO(base64.b64decode(_image) if isinstance(_image, str) else _image))
    elif isinstance(_image, Image.Image):
        img = _image
    else:
        img = Image.open(_image)

    # Konvertiere in RGB-Modus, falls erforderlich
    if img.mode != 'RGB':
        img = img.convert('RGB')

    # Verkleinern, wenn das Bild zu groß ist
    max_size = 1000  # Reduzierte Maximalgröße zur Verringerung des Speicherverbrauchs
    if max(img.size) > max_size:
        img.thumbnail((max_size, max_size))

    # Speichern in Bytes
    img_byte_arr = io.BytesIO()
    img.save(img_byte_arr, format='JPEG')
    img_byte_arr = img_byte_arr.getvalue()

    return base64.b64encode(img_byte_arr).decode('utf-8')

def replace_german_sharp_s(text):
    """Ersetzt alle Vorkommen von 'ß' durch 'ss'."""
    return text.replace('ß', 'ss')

def clean_json_string(s):
    s = s.strip()
    s = re.sub(r'^json\s*', '', s)
    s = re.sub(r'\s*$', '', s)
    s = re.sub(r'\s+', ' ', s)
    s = re.sub(r'(?<=text": ")(.+?)(?=")', lambda m: m.group(1).replace('\n', '\\n'), s)
    s = ''.join(char for char in s if ord(char) >= 32 or char == '\n')
    match = re.search(r'\[.*\]', s, re.DOTALL)
    return match.group(0) if match else s

def convert_json_to_text_format(json_input):
    if isinstance(json_input, str):
        data = json.loads(json_input)
    else:
        data = json_input

    fib_output = []
    ic_output = []

    for item in data:
        text = item.get('text', '')
        blanks = item.get('blanks', [])
        wrong_substitutes = item.get('wrong_substitutes', [])

        num_blanks = len(blanks)

        fib_lines = [
            "Type\tFIB",
            "Title\t✏✏Vervollständigen Sie die Lücken mit dem korrekten Begriff.✏✏",
            f"Points\t{num_blanks}"
        ]

        for blank in blanks:
            text = text.replace(blank, "{blank}", 1)

        parts = text.split("{blank}")
        for index, part in enumerate(parts):
            fib_lines.append(f"Text\t{part.strip()}")
            if index < len(blanks):
                fib_lines.append(f"1\t{blanks[index]}\t20")

        fib_output.append('\n'.join(fib_lines))

        ic_lines = [
            "Type\tInlinechoice",
            "Title\tWörter einordnen",
            "Question\t✏✏Wählen Sie die richtigen Wörter.✏✏",
            f"Points\t{num_blanks}"
        ]

        all_options = blanks + wrong_substitutes
        random.shuffle(all_options)

        for index, part in enumerate(parts):
            ic_lines.append(f"Text\t{part.strip()}")
            if index < len(blanks):
                options_str = '|'.join(all_options)
                ic_lines.append(f"1\t{options_str}\t{blanks[index]}\t|")

        ic_output.append('\n'.join(ic_lines))

    return '\n\n'.join(fib_output), '\n\n'.join(ic_output)

def transform_output(json_string):
    try:
        cleaned_json_string = clean_json_string(json_string)
        json_data = json.loads(cleaned_json_string)
        fib_output, ic_output = convert_json_to_text_format(json_data)
        
        # Anwenden der Reinigungsfunktion
        fib_output = replace_german_sharp_s(fib_output)
        ic_output = replace_german_sharp_s(ic_output)

        return f"{ic_output}\n---\n{fib_output}"
    except json.JSONDecodeError as e:
        st.error(f"Fehler beim Parsen von JSON: {e}")
        st.text("Bereinigte Eingabe:")
        st.code(cleaned_json_string, language='json')
        st.text("Originale Eingabe:")
        st.code(json_string)
        
        try:
            if not cleaned_json_string.strip().endswith(']'):
                cleaned_json_string += ']'
            partial_json = json.loads(cleaned_json_string)
            st.warning("Teilweises JSON konnte gerettet werden. Ergebnisse können unvollständig sein.")
            fib_output, ic_output = convert_json_to_text_format(partial_json)
            return f"{ic_output}\n---\n{fib_output}"
        except:
            st.error("Teilweises JSON konnte nicht gerettet werden.")
            return "Fehler: Ungültiges JSON-Format"
    except Exception as e:
        st.error(f"Fehler bei der Verarbeitung der Eingabe: {str(e)}")
        st.text("Originale Eingabe:")
        st.code(json_string)
        return "Fehler: Eingabe konnte nicht verarbeitet werden"

def get_chatgpt_response(prompt, model, image=None, selected_language="English"):
    """Ruft eine Antwort von OpenAI GPT mit Fehlerbehandlung ab."""
    try:
        # System-Prompt erstellen, der Sprachinstruktionen enthält
        system_prompt = (
            """
            Du bist ein Experte im Bildungsbereich, spezialisiert auf die Erstellung von Testfragen und -antworten zu allen Themen, unter Einhaltung der Bloom's Taxonomy. Deine Aufgabe ist es, hochwertige Frage-Antwort-Sets basierend auf dem vom Benutzer bereitgestellten Material zu erstellen, wobei jede Frage einer spezifischen Ebene der Bloom's Taxonomy entspricht: Erinnern, Verstehen, Anwenden, Analysieren, Bewerten und Erstellen.

            Der Benutzer wird entweder Text oder ein Bild hochladen. Deine Aufgaben sind wie folgt:

            **Input-Analyse:**

            - Du analysierst du den Inhalt sorgfältig, um die Schlüsselkonzepte und wichtigen Informationen zu verstehen.
            - Falls vorhanden, du achtest auf Diagramme, Grafiken, Bilder oder Infografiken, um Bildungsinhalte abzuleiten.

            **Fragen-Generierung nach Bloom-Ebene:**
            Basierend auf dem analysierten Material generierst du Fragen über alle die folgenden Ebenen der Bloom's Taxonomy:

            - **Erinnern**: Einfache, abrufbasierte Fragen.
            - **Verstehen**: Fragen, die das Verständnis des Materials bewerten.
            - **Anwenden**: Fragen, die die Anwendung des Wissens in praktischen Situationen erfordern.
            """
        )
        
        if image:
            base64_image = process_image(image)
            messages = [
                {"role": "system", "content": system_prompt},
                {
                    "role": "user", 
                    "content": [
                        {"type": "text", "text": prompt},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{base64_image}",
                                "detail": "low"
                            }
                        }
                    ]
                }
            ]
        else:
            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": prompt}
            ]

        response = client.chat.completions.create(
            model=model,
            messages=messages,
            max_tokens=15000,  # Updated max tokens
            temperature=0.6
        )
        
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"Fehler bei der Kommunikation mit der OpenAI API: {e}")
        logging.error(f"Fehler bei der Kommunikation mit der OpenAI API: {e}")
        return None

def process_images(images, selected_language, selected_model):
    """Verarbeitet hochgeladene Bilder und generiert Fragen."""
    for idx, image in enumerate(images):
        st.image(image, caption=f'Seite {idx+1}', use_column_width=True)

        # Textbereich für Benutzereingaben und Lernziele
        user_input = st.text_area(f"Geben Sie Ihre Frage oder Anweisungen für Seite {idx+1} ein:", key=f"text_area_{idx}")
        learning_goals = st.text_area(f"Lernziele für Seite {idx+1} (Optional):", key=f"learning_goals_{idx}")
        selected_types = st.multiselect(f"Wählen Sie die Fragetypen für Seite {idx+1} aus:", MESSAGE_TYPES, key=f"selected_types_{idx}")

        # Button zum Generieren von Fragen für die Seite
        if st.button(f"Fragen für Seite {idx+1} generieren", key=f"generate_button_{idx}"):
            # Fragen nur generieren, wenn Benutzereingaben und ausgewählte Fragetypen vorhanden sind
            if user_input and selected_types:
                # Übergabe der ausgewählten Sprache und des Modells hier
                generate_questions_with_image(user_input, learning_goals, selected_types, image, selected_language, selected_model)
            else:
                st.warning(f"Bitte geben Sie Text ein und wählen Sie Fragetypen für Seite {idx+1} aus.")

def generate_questions_with_image(user_input, learning_goals, selected_types, image, selected_language, selected_model):
    """Generiert Fragen für das Bild und behandelt Fehler."""
    all_responses = ""
    generated_content = {}
    for msg_type in selected_types:
        prompt_template = read_prompt_from_md(msg_type)
        full_prompt = f"{prompt_template}\n\nBenutzereingabe: {user_input}\n\nLernziele: {learning_goals}"
        try:
            response = get_chatgpt_response(full_prompt, model=selected_model, image=image, selected_language=selected_language)
            if response:
                if msg_type == "inline_fib":
                    processed_response = transform_output(response)
                    generated_content[f"{msg_type.replace('_', ' ').title()} (Verarbeitet)"] = processed_response
                    all_responses += f"{processed_response}\n\n"
                else:
                    generated_content[msg_type.replace('_', ' ').title()] = response
                    all_responses += f"{response}\n\n"
            else:
                st.error(f"Fehler bei der Generierung einer Antwort für {msg_type}.")
        except Exception as e:
            st.error(f"Ein Fehler ist für {msg_type} aufgetreten: {str(e)}")
    
    # Reinigungsfunktion auf alle Antworten anwenden
    all_responses = replace_german_sharp_s(all_responses)

    # Generierten Inhalt mit Häkchen anzeigen
    st.subheader("Generierter Inhalt:")
    for title in generated_content.keys():
        st.write(f"✔ {title}")

    # Download-Button für alle Antworten
    if all_responses:
        st.download_button(
            label="Alle Antworten herunterladen",
            data=all_responses,
            file_name="alle_antworten.txt",
            mime="text/plain"
        )

@st.cache_data
def convert_pdf_to_images(file):
    """Konvertiert PDF-Seiten in Bilder."""
    images = convert_from_bytes(file.read())
    return images

@st.cache_data
def extract_text_from_pdf(file):
    """Extrahiert Text aus PDF mit PyPDF2."""
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text
    return text.strip()

@st.cache_data
def extract_text_from_docx(file):
    """Extrahiert Text aus DOCX-Datei."""
    doc = docx.Document(file)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return text.strip()

def is_pdf_ocr(text):
    """Prüft, ob das PDF OCR-Text enthält (Implementierung erforderlich)."""
    # Dummy-Implementierung, bitte nach Bedarf anpassen
    return bool(text)

def process_pdf(file):
    text_content = extract_text_from_pdf(file)
    
    # Wenn kein Text gefunden wurde, nehme an, dass es ein nicht-OCR-PDF ist
    if not text_content or not is_pdf_ocr(text_content):
        st.warning("Dieses PDF ist nicht OCR-geschützt. Textextraktion fehlgeschlagen. Bitte laden Sie ein OCR-PDF hoch.")
        return None, convert_pdf_to_images(file)  # Fallback zur Bildverarbeitung
    else:
        return text_content, None

def main():
    """Hauptfunktion für die Streamlit-App."""
    # Modellenauswahl mit Dropdown
    st.subheader("Modell für die Generierung auswählen:")
    model_options = ["gpt-4o", "gpt-4o-mini"]
    selected_model = st.selectbox("Wählen Sie das Modell aus:", model_options, index=0)

    # Sprachenauswahl mit Radiobuttons
    st.subheader("Sprache für generierte Fragen auswählen:")
    languages = {
        "Deutsch": "German",
        "Englisch": "English",
        "Französisch": "French",
        "Italienisch": "Italian",
        "Spanisch": "Spanish"
    }
    selected_language = st.radio("Wählen Sie die Sprache für die Ausgabe:", list(languages.keys()), index=0)

    # Dateiuploader-Bereich
    uploaded_file = st.file_uploader("Laden Sie eine PDF, DOCX oder Bilddatei hoch", type=["pdf", "docx", "jpg", "jpeg", "png"])

    text_content = ""
    image_content = None
    images = []

    # Cache zurücksetzen, wenn eine neue Datei hochgeladen wird
    if uploaded_file:
        st.cache_data.clear()  # Dies löscht den Cache, um vorherige zwischengespeicherte Werte zu vermeiden

    if uploaded_file is not None:
        if uploaded_file.type == "application/pdf":
            text_content, images = process_pdf(uploaded_file)
            if text_content:
                st.success("Text aus PDF extrahiert. Sie können ihn nun im folgenden Textfeld bearbeiten. PDFs, die länger als 5 Seiten sind, sollten gekürzt werden.")
            elif images:
                st.success("PDF in Bilder konvertiert. Sie können jetzt Fragen zu jeder Seite stellen.")
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text_content = extract_text_from_docx(uploaded_file)
            st.success("Text erfolgreich extrahiert. Sie können ihn nun im folgenden Textbereich bearbeiten.")
        elif uploaded_file.type.startswith('image/'):
            image_content = Image.open(uploaded_file)
            st.image(image_content, caption='Hochgeladenes Bild', use_column_width=True)
            st.success("Bild erfolgreich hochgeladen. Sie können jetzt Fragen zum Bild stellen.")
        else:
            st.error("Nicht unterstützter Dateityp. Bitte laden Sie eine PDF, DOCX oder Bilddatei hoch.")

    # Bilder verarbeiten, falls vorhanden, ansonsten Text oder Bildinhalt verarbeiten
    if images:
        process_images(images, selected_language, selected_model)  # Übergabe der ausgewählten Sprache und des Modells hier
    else:
        user_input = st.text_area("Geben Sie Ihren Text oder Ihre Frage zum Bild ein:", value=text_content)
        learning_goals = st.text_area("Lernziele (Optional):")
        
        # Fragetypen auswählen
        selected_types = st.multiselect("Wählen Sie die Fragetypen zur Generierung aus:", MESSAGE_TYPES)
        
        # Benutzerdefiniertes CSS für hellblauen Hintergrund in Info-Callouts
        st.markdown(
            """
            <style>
            .custom-info {
                background-color: #e7f3fe;
                padding: 10px;
                border-radius: 5px;
                border-left: 6px solid #2196F3;
            }
            .custom-success {
                background-color: #d4edda;
                padding: 10px;
                border-radius: 5px;
                border-left: 6px solid #28a745;
            }
            .custom-warning {
                background-color: #fff3cd;
                padding: 10px;
                border-radius: 5px;
                border-left: 6px solid #ffc107;
            }
            </style>
            """, unsafe_allow_html=True
        )

        # Button zum Generieren von Fragen
        if st.button("Fragen generieren"):
            if (user_input or image_content) and selected_types:
                generate_questions_with_image(user_input, learning_goals, selected_types, image_content, selected_language, selected_model)              
            elif not user_input and not image_content:
                st.warning("Bitte geben Sie etwas Text ein, laden Sie eine Datei hoch oder laden Sie ein Bild hoch.")
            elif not selected_types:
                st.warning("Bitte wählen Sie mindestens einen Fragetyp aus.")

if __name__ == "__main__":
    main()
